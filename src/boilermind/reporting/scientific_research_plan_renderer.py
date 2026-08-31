"""Deterministic Markdown and Word rendering for the scientific plan."""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from boilermind.core.contracts.scientific_research_plan import ScientificResearchPlan


def _text(value: Any) -> str:
    if value is None or value == "":
        return "未提供"
    if isinstance(value, (dict, list, tuple)):
        value = json.dumps(value, ensure_ascii=False, sort_keys=True)
    return _localize_report_text(str(value))


def _localize_report_text(text: str) -> str:
    if text.casefold() == "supported":
        return "得到支持"
    labels = {
        "all_candidates_worse_than_reference_on:MAE": "所有候选模型在 MAE 上均劣于参考模型",
        "any_candidate_better_than_reference_on:MAE": "任一候选模型在 MAE 上优于参考模型",
        "locked_test_not_used_for_selection": "锁定测试集不参与模型选择",
        "locked_test": "锁定测试集",
        "validation_only_model_selection": "仅使用验证集进行模型选择",
        "reference_model_comparison": "参考基线对比",
        "chronological_validation": "按时间顺序验证",
        "locked_test_evaluation": "锁定测试集评估",
        "model_comparison": "模型对比",
        "UNIFIED_EXPERIMENT_EXECUTION": "统一实验执行",
        "target_profile:": "目标变量：",
        "compare_prediction_accuracy": "比较预测精度",
        "unspecified": "未特别限定",
        "falsified": "未获支持",
        "partially_supported": "部分支持",
        "insufficient_evidence": "证据不足",
        "chronological validation + locked test": "按时间顺序验证 + 锁定测试集",
        "train/validation/locked-test": "训练集/验证集/锁定测试集",
        "ExperimentAudit": "实验审计",
        "ResearchProblemSpec": "研究问题规范",
        "EvidenceBundle": "证据集合",
        "ScientificHypothesis": "科学假设",
        "ExperimentPlan": "实验方案",
        "ExperimentContract": "实验契约",
        "ExperimentResult": "实验结果",
        "ScientificResult": "科学结果",
        "FIRST_ROUND_NO_ITERATION": "首轮无需迭代",
        "COMPLETED": "已完成",
        "data_preprocessing": "数据预处理",
        "evaluation": "评估",
        "audit": "审计",
        "baseline": "基线",
        "current_operating_point": "当前工况点",
        "PROPOSED_COMPETING_HYPOTHESIS": "待验证",
        "Persistence": "持久性基线模型",
        "ACHIEVED_RISE_PCT": "目标提升比例（%）",
        "PRESSURE_MAX_MPA": "最大压力（MPa）",
        "hgb_control_optimizer": "HGB 控制优化器",
        "hgb_soft_sensor_validation": "HGB 软测验证",
        "constrained_control_optimization": "受约束控制优化",
        "drum_pressure <= 23 MPa": "汽包压力不超过 23 MPa",
        "each control adjustment within ±25%": "各控制变量的单次调整幅度不超过 ±25%",
        "constraint_search_from_validation_operating_point": "基于验证集工况点的约束搜索",
        "constraint_search": "约束搜索",
        "unity_result_export": "Unity 结果导出",
        "steam_volumetric_flow": "蒸汽体积流量 V",
        "POST_EXPERIMENT": "实验后",
        "POST_ITERATION": "迭代后",
        "ACTUAL_EXECUTION": "实际执行结果",
        "validation_and_locked_test": "验证集与锁定测试集",
        "validation_only_model_selection": "仅使用验证集进行模型选择",
        "chronological validation": "按时间顺序验证",
        "chronological locked test": "按时间顺序锁定测试",
        "chronological": "按时间顺序",
        "test_frozen": "锁定测试集",
        "locked test 最优": "锁定测试集最优",
        "使用 locked test 回选": "使用锁定测试集回选",
        "locked test 不参与": "锁定测试集不参与",
        "locked-test": "锁定测试集",
        "locked test": "锁定测试集",
        "validation-only": "仅验证集",
        "validation": "验证集",
        "train": "训练集",
        "VERIFIED": "已核验",
        "True": "是",
        "False": "否",
        "random_state": "随机种子",
        "max_depth": "最大深度",
        "max_iter": "最大迭代次数",
        "learning_rate": "学习率",
        "None": "未提供",
    }
    for source, target in labels.items():
        text = text.replace(source, target)
    text = re.sub(r"\bsupported\b", "得到支持", text, flags=re.IGNORECASE)
    text = re.sub(r"\bfalsified\b", "未获支持", text, flags=re.IGNORECASE)
    text = re.sub(r"\bgru\b", "GRU", text, flags=re.IGNORECASE)
    text = re.sub(r"\blstm\b", "LSTM", text, flags=re.IGNORECASE)
    return re.sub(r"\bcompleted\b", "已完成", text, flags=re.IGNORECASE)


def _bullets(values: list[Any] | tuple[Any, ...] | None) -> list[str]:
    return [f"- {_text(value)}" for value in (values or [])] or ["- 无"]


def _sub_bullets(values: list[Any] | tuple[Any, ...] | None) -> list[str]:
    return [f"  - {_text(value)}" for value in (values or [])] or ["  - 无"]


def _number(value: Any) -> str:
    try:
        return f"{float(value):.6f}"
    except (TypeError, ValueError):
        return "—"


def _control_number(value: Any) -> str:
    try:
        number = float(value)
    except (TypeError, ValueError):
        return "—"
    text = f"{number:.4f}".rstrip("0").rstrip(".")
    return text or "0"


def _metric(metrics: dict[str, Any] | None, name: str) -> Any:
    wanted = name.casefold()
    for key, value in (metrics or {}).items():
        if str(key).casefold() == wanted:
            return value
    return None


def _model_name(name: str | None) -> str:
    labels = {
        "gru": "GRU",
        "lstm": "LSTM",
        "rf": "RandomForest",
        "bayesianridge": "BayesianRidge",
        "ridge": "Ridge",
        "persistence": "持久性基线模型",
        "hgb_control_optimizer": "HGB 控制优化器",
    }
    return labels.get(str(name or "").casefold(), str(name or "未形成"))


class ScientificResearchPlanRenderer:
    """Render only frozen report fields; never infer new scientific claims."""

    def render_markdown(self, plan: ScientificResearchPlan) -> str:
        abstract = plan.paper_abstract
        problem, rationale = plan.problem_statement, plan.rationale
        technical, dataset = plan.technical_details, plan.dataset
        methods, experiments, results = plan.methods, plan.experiments, plan.results
        metrics = plan.metrics
        validity, verdict = plan.experiment_validity, plan.scientific_verdict
        references = plan.references or []
        competing = rationale.competing_hypotheses if rationale else []
        metric_names = [
            name for name in [
                metrics.primary_metric if metrics else None,
                *(metrics.secondary_metrics if metrics else []),
            ] if name
        ]
        validation = metrics.validation_metrics_by_model if metrics else {}
        locked = metrics.locked_test_metrics_by_model if metrics else {}
        models = list(dict.fromkeys([*(validation or {}), *(locked or {})]))
        baseline = metrics.baseline_metrics if metrics else {}
        abstract_text = (
            getattr(abstract, "rendered_text", None)
            or getattr(abstract, "polished_text", None)
            or _text(abstract)
        )
        lines = [
            "# 《科研假设与研究计划》", "",
            f"## {_text(plan.paper_title)}", "",
            "> 本报告基于本次实验的数据、方法和结果整理而成。模型选择依据验证集，"
            "锁定测试集仅用于独立检验。", "",
            "## 1. 研究摘要", "",
            f"{_text(abstract_text)}", "",
            "### 核心结论", "",
            f"- 本次选定模型：**{_model_name(results.protocol_selected_model if results else None)}**",
            f"- 锁定测试集表现最佳：**{_model_name(results.locked_test_best_model if results else None)}**",
            f"- 科学裁决：**{_text(results.scientific_verdict if results else None)}**",
            f"- 解释：{_text(results.selection_interpretation if results else None)}", "",
            "## 2. 科研问题与研究边界", "",
            f"- 原始问题：{_text(problem.original_question if problem else None)}",
            f"- 研究对象：{_text(problem.research_object if problem else None)}",
            f"- 目标变量：{_text(problem.target_variable if problem else None)}",
            f"- 预测时域：{('软测（当前时刻）' if (technical and technical.prediction_horizon_steps is None) else f'{_text(technical.prediction_horizon_steps if technical else None)} 步')}",
            f"- 工况范围：{_text(problem.operating_condition if problem else None)}",
            f"- 研究缺口：{_text(problem.research_gap if problem else None)}", "",
            "## 3. 主假设与竞争子假设", "",
            f"### 3.1 主假设", "",
            _text(rationale.hypothesis_statement if rationale else None), "",
            f"- 假设创新点：{_text(rationale.innovation_point if rationale else None)}",
            f"- 机理推理链：{_text(rationale.mechanism_chain if rationale else None)}",
            "",
            "### 3.2 竞争子假设", "",
        ]
        for index, item in enumerate(competing, start=1):
            experiment_plan = item.get("experiment_plan") or {}
            lines.extend([
                f"#### H{index} · {_text(item.get('title'))}（{_model_name(item.get('model'))}）",
                "",
                f"- 假设：{_text(item.get('statement'))}",
                f"- 预期观察：{_text(item.get('expected_observation'))}",
                f"- 状态：{_text(item.get('status'))}", "",
                "**验证设计**", "",
                f"- 研究目的：{_text(experiment_plan.get('objective'))}",
                f"- 比较方法：{_text(experiment_plan.get('design'))}",
                f"- 评价指标：{_text(experiment_plan.get('primary_endpoint'))}",
                f"- 确认规则：{_text(experiment_plan.get('confirmation_rule'))}",
                f"- 证伪规则：{_text(experiment_plan.get('falsification_rule'))}",
                f"- 当前进展：{_text(experiment_plan.get('execution_status'))}", "",
            ])
        lines.extend([
            "### 3.3 确认、证伪与反机制", "",
            "- 确认标准：", *_sub_bullets(rationale.confirmation_criteria if rationale else []),
            "- 证伪标准：", *_sub_bullets(rationale.falsification_criteria if rationale else []),
            "- 关键假设条件：", *_sub_bullets(rationale.assumptions if rationale else []), "",
            "## 4. 研究设计", "",
            "### 4.1 数据与切分", "",
            f"- 数据集：{_text(dataset.dataset_id if dataset else None)}",
            f"- 输入特征数：{len(dataset.input_variables or []) if dataset else 0}",
            f"- 历史窗口：{_text(technical.window_steps if technical else None)} 步",
            f"- 训练/验证/锁定测试：{_text(dataset.train_split if dataset else None)} / "
            f"{_text(dataset.validation_split if dataset else None)} / "
            f"{_text(dataset.locked_test_split if dataset else None)}",
            "- 预处理：仅允许在训练集上拟合。",
            "- 选模：仅使用验证集；锁定测试集冻结后一次性评估。", "",
            "### 4.2 实施步骤", "",
        ])
        for item in (methods.implementation_steps if methods else []):
            lines.append(
                f"{item.get('step_index', '')}. **{_text(item.get('name'))}**："
                f"{_text(item.get('description'))}"
            )
        control = methods.control if methods else None
        treatment = methods.treatment if methods else None
        if control or treatment:
            lines.extend(["", "### 4.3 控制优化方案", ""])
            if isinstance(control, dict):
                current_values = control.get("current_values") or {}
                current_volume = control.get("current_volume")
                lines.append("- 当前工况：")
                for name, value in current_values.items():
                    lines.append(f"  - {name}：{_control_number(value)}")
                if current_volume is not None:
                    lines.append(f"  - 当前蒸汽体积量 V：{_control_number(current_volume)}")
            if isinstance(treatment, dict):
                ranges = treatment.get("adjustment_ranges") or {}
                recommended = treatment.get("recommended_values") or {}
                if ranges:
                    lines.extend([
                        "",
                        "| 变量 | 当前值 | 建议范围下限 | 建议范围上限 | 推荐值 |",
                        "|---|--:|--:|--:|--:|",
                    ])
                    for name, item in ranges.items():
                        if isinstance(item, dict):
                            lines.append(
                                f"| {name} | {_control_number(item.get('current'))} | "
                                f"{_control_number(item.get('minimum'))} | "
                                f"{_control_number(item.get('maximum'))} | "
                                f"{_control_number(item.get('recommended'))} |"
                            )
                        else:
                            lines.append(
                                f"| {name} | — | — | — | {_text(item)} |"
                            )
                if recommended and not ranges:
                    lines.extend(["", "- 推荐单点控制值："])
                    for name, value in recommended.items():
                        lines.append(f"  - {name}：{_control_number(value)}")
            if problem and problem.constraints:
                lines.extend(["", "- 运行约束：", *_sub_bullets(problem.constraints)])
            if methods and methods.objective:
                lines.extend(["", f"- 优化目标：{_text(methods.objective)}"])
        lines.extend([
            "", "## 5. 模型比较结果", "",
            "### 5.1 验证集（用于选择）", "",
            "| 模型 | " + " | ".join(_text(name) for name in metric_names) + " |",
            "|---|" + "---:|" * len(metric_names),
        ])
        for model in models:
            values = [_number(_metric((validation or {}).get(model), name)) for name in metric_names]
            marker = " **← 协议选择**" if results and model == results.protocol_selected_model else ""
            lines.append(f"| {_model_name(model)}{marker} | " + " | ".join(values) + " |")
        lines.extend([
            "", "### 5.2 锁定测试集（仅用于泛化评价）", "",
            "| 模型 | " + " | ".join(_text(name) for name in metric_names) + " |",
            "|---|" + "---:|" * len(metric_names),
        ])
        for model in models:
            values = [_number(_metric((locked or {}).get(model), name)) for name in metric_names]
            marker = " **← 测试集最低误差**" if results and model == results.locked_test_best_model else ""
            lines.append(f"| {_model_name(model)}{marker} | " + " | ".join(values) + " |")
        if baseline:
            values = [_number(_metric(baseline, name)) for name in metric_names]
            lines.append("| 持久性基线模型（参考） | " + " | ".join(values) + " |")
        lines.extend([
            "", "### 5.3 科学解释", "",
            _text(results.selection_interpretation if results else None), "",
            "正式选择与锁定测试集排名不一致并非允许改选模型的理由，而是提示时间分布漂移、"
            "模型方差或工况组成变化，需要通过多时间块、多随机种子和工况分层复验。", "",
        ])
        if experiments and (experiments.execution_notes or experiments.artifacts):
            lines.extend(["### 5.4 执行情况", ""])
            if experiments.execution_notes:
                lines.extend(["- 执行记录：", *_sub_bullets(experiments.execution_notes)])
            lines.append("")
        lines.extend([
            "## 6. 研究结论与验证情况", "",
            f"- 研究结论：**{_text(results.scientific_verdict if results else None)}**",
            "- 已满足标准：", *_sub_bullets(results.achieved_criteria if results else []),
            "- 未满足标准：", *_sub_bullets(results.failed_criteria if results else []),
            f"- 执行有效：{_text(validity.execution_valid if validity else None)}",
            f"- 数据冻结：{_text(validity.dataset_frozen if validity else None)}",
            f"- 泄漏检查：{_text(validity.leakage_check_passed if validity else None)}",
            f"- 基线有效：{_text(validity.baseline_valid if validity else None)}",
            f"- 指标检查：{_text(validity.metric_check_passed if validity else None)}", "",
            "## 7. 下一轮研究建议", "",
            "1. 对协议选择模型开展多随机种子复验，报告均值、标准差和最差值。",
            "2. 使用多个冻结时间块检验排名是否稳定，避免单一测试时间段偶然性。",
            "3. 按稳态、升负荷、降负荷和方向变化进行工况分层评价。",
            "4. 复核目标变量单位和指标单位的一致性，再进入工程部署评审。",
            "5. 若验证集与锁定测试集排名持续不一致，重新审查数据漂移与模型选择策略。", "",
        ])
        if technical and "control" in str(technical.experiment_type or ""):
            lines.extend([
                "6. Unity 后续验证：将推荐控制值与调参范围推送 Unity，回传实际蒸汽体积量，"
                "对比小模型预测偏差并形成第二层科学裁决（支持/部分支持/证据不足/证伪）。",
                "",
            ])
        lines.extend([
            "## 8. 参考文献", "",
            "著录格式：GB/T 7714—2015（顺序编码制）。", "",
        ])
        if references:
            for index, ref in enumerate(references, start=1):
                lines.append(
                    f"[{index}] {_text(ref.formatted_citation or ref.citation or ref.title)}"
                )
        else:
            lines.append("无可验证参考论文；系统未补造文献。")
        lines.extend([
            "", "## 9. 局限性与适用范围", "",
            *_bullets(plan.limitations),
        ])
        return "\n".join(lines).rstrip() + "\n"

    def render_docx(self, plan: ScientificResearchPlan, path: str | Path) -> Path:
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Inches, Pt, RGBColor
        except ImportError as exc:
            raise RuntimeError("python-docx dependency is required for Word output") from exc
        target = Path(path)
        target.parent.mkdir(parents=True, exist_ok=True)
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        styles = document.styles
        styles["Normal"].font.name = "Microsoft YaHei"
        styles["Normal"].font.size = Pt(10.5)
        title = document.add_heading("科研假设与研究计划", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle = document.add_paragraph(_text(plan.paper_title))
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(14)
        subtitle.runs[0].font.color.rgb = RGBColor(55, 65, 81)
        markdown = self.render_markdown(plan)
        lines = markdown.splitlines()[4:]
        index = 0
        while index < len(lines):
            line = lines[index]
            if line.startswith("## "):
                document.add_heading(line[3:], level=1)
            elif line.startswith("### "):
                document.add_heading(line[4:], level=2)
            elif line.startswith("#### "):
                document.add_heading(line[5:], level=3)
            elif line.startswith("  - "):
                document.add_paragraph(line[4:].replace("**", "").replace("`", ""), style="List Bullet 2")
            elif line.startswith("  ") and line.strip()[0].isdigit() and ". " in line.strip()[:5]:
                document.add_paragraph(line.strip().replace("**", "").replace("`", ""))
            elif line.startswith("- "):
                document.add_paragraph(line[2:].replace("**", "").replace("`", ""), style="List Bullet")
            elif line.startswith("> "):
                paragraph = document.add_paragraph(line[2:])
                paragraph.style = styles["Quote"]
            elif line.startswith("|") and index + 1 < len(lines) and lines[index + 1].startswith("|---"):
                headers = [cell.strip() for cell in line.strip("|").split("|")]
                index += 2
                rows = []
                while index < len(lines) and lines[index].startswith("|"):
                    rows.append([cell.strip() for cell in lines[index].strip("|").split("|")])
                    index += 1
                table = document.add_table(rows=1, cols=len(headers))
                table.style = "Table Grid"
                for column, value in enumerate(headers):
                    table.rows[0].cells[column].text = value
                for values in rows:
                    cells = table.add_row().cells
                    for column, value in enumerate(values[:len(cells)]):
                        cells[column].text = value.replace("**", "")
                index -= 1
            elif line and line[0].isdigit() and ". " in line[:5]:
                document.add_paragraph(line.replace("**", "").replace("`", ""))
            elif line:
                document.add_paragraph(line.replace("**", "").replace("`", ""))
            index += 1
        footer = section.footer.paragraphs[0]
        footer.text = (
            f"运行编号 {plan.metadata.run_id if plan.metadata else '未提供'} · "
            "BoilerMind 可审计科研输出"
        )
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.save(target)
        return target
